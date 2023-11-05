from libretranslatepy import LibreTranslateAPI
from sacremoses import MosesTokenizer, MosesDetokenizer
import docx

lt = LibreTranslateAPI("https://translate.argosopentech.com/")  # Use your preferred LibreTranslate API endpoint

# Define input and output languages
input_lang = 'en'
output_languages = ["fr", "de", "ru"]  # French, German, Russian

# Initialize Moses tokenizer and detokenizer
tokenizer = MosesTokenizer(lang=input_lang)
detokenizer = MosesDetokenizer(lang=input_lang)

# Read text from a file (you can adapt this to your input source)
with open("input.txt", "r", encoding="utf-8") as file:
    text = file.read()

# Translate text into multiple languages and store results
translated_texts = {}

for lang in output_languages:
    # Tokenize the text
    tokenized_text = tokenizer.tokenize(text, return_str=True)
    
    # Translate the tokenized text
    translated_text = lt.translate(tokenized_text, input_lang, lang)
    
    # Detokenize the translated text
    detokenized_text = detokenizer.detokenize(translated_text.split('\n'))

    translated_texts[lang] = detokenized_text

# Create a DOCX document for each language
for lang, translated_text in translated_texts.items():
    doc = docx.Document()
    doc.add_heading(f"Translation to {lang}:", level=1)
    doc.add_paragraph(translated_text)

    # Save the DOCX document
    doc.save(f"translated_text_{lang}.docx")

print("Translation to DOCX Successful")
