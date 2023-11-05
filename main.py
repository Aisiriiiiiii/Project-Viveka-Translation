import sys
from PyQt5.QtWidgets import *
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5 import uic
from PyQt5.uic import loadUi
from PyQt5.QtCore import QUrl, Qt
from PyQt5.QtGui import QDesktopServices
from PyQt5.QtWidgets import QMainWindow, QApplication, QMessageBox
from sacremoses import MosesTokenizer, MosesDetokenizer
import docx
from libretranslatepy import LibreTranslateAPI
lt = LibreTranslateAPI("https://translate.argosopentech.com/") 
input_lang = 'en'
fname = ''
docpath = ''


class MyGUI(QMainWindow):
    
    def __init__(self):
            super(MyGUI,self).__init__()
            uic.loadUi('mw.ui',self)
            self.show()
            self.pushButton.clicked.connect(self.browsefiles)
            self.pushButton_2.clicked.connect(self.openwindow2)
            
                        
    def browsefiles(self):
         global fname
         fname,_ =QFileDialog.getOpenFileName(self, 'Open file')
         self.lineEdit.setText(fname)      
         
    def openwindow2(self):
        self.window = QtWidgets.QMainWindow()
        self.ui = SecondWindow()
        self.window.show()
        
        
        
class SecondWindow(QMainWindow):
    def __init__(self):
        super(SecondWindow,self).__init__()
        uic.loadUi('secondwindow.ui',self)
        self.german.clicked.connect(self.translateToGerman)
        self.french.clicked.connect(self.translateToFrench)
        self.russian.clicked.connect(self.translateToRussian)
        self.show()
    global docpath
    
    def translateToGerman(self):
        global fname
        tokenizer = MosesTokenizer(lang='en')
        detokenizer = MosesDetokenizer(lang='en')
        with open(fname, "r", encoding="utf-8") as file:
                text = file.read()
        translated_texts = {}
        tokenized_text = tokenizer.tokenize(text, return_str=True)
        translated_text = lt.translate(tokenized_text, 'en', 'ge')
        detokenized_text = detokenizer.detokenize(translated_text.split('\n'))
        translated_texts[lang] = detokenized_text


        for lang, translated_text in translated_texts.items():
            doc = docx.Document()
            doc.add_heading(f"Translation to {lang}:", level=1)
            doc.add_paragraph(translated_text)
            docpath = doc.save(f"translated_text_{lang}.docx")

        print("Translation to DOCX Successful") 
        
        self.display_window = FinalWindow()
        self.display_window.show()       

    def translateToFrench(self):
        global fname
        tokenizer = MosesTokenizer(lang='en')
        detokenizer = MosesDetokenizer(lang='en')
        
        with open(fname, "r", encoding="utf-8") as file:
                text = file.read()

        translated_texts = {}
        tokenized_text = tokenizer.tokenize(text, return_str=True)
        translated_text = lt.translate(tokenized_text, 'en', 'fr')
        detokenized_text = detokenizer.detokenize(translated_text.split('\n'))
        translated_texts[lang] = detokenized_text


        for lang, translated_text in translated_texts.items():
            doc = docx.Document()
            doc.add_heading(f"Translation to {lang}:", level=1)
            doc.add_paragraph(translated_text)
            docpath = doc.save(f"translated_text_{lang}.docx")
            
        print("Translation to DOCX Successful") 
        
        self.display_window = FinalWindow()
        self.display_window.show()  
        
        
    def translateToRussian(self):
        global fname
        tokenizer = MosesTokenizer(lang='en')
        detokenizer = MosesDetokenizer(lang='en')
        
        with open(fname, "r", encoding="utf-8") as file:
                text = file.read()

        translated_texts = {}
        tokenized_text = tokenizer.tokenize(text, return_str=True)
        translated_text = lt.translate(tokenized_text, 'en', 'ru')
        detokenized_text = detokenizer.detokenize(translated_text.split('\n'))
        translated_texts[lang] = detokenized_text


        for lang, translated_text in translated_texts.items():
            doc = docx.Document()
            doc.add_heading(f"Translation to {lang}:", level=1)
            doc.add_paragraph(translated_text)
            docpath = doc.save(f"translated_text_{lang}.docx") 
            
            
        print("Translation to DOCX Successful")
        
        self.display_window = FinalWindow()
        self.display_window.show()
            
        
        
class FinalWindow(QMainWindow):
    def __init__(self):
        super(FinalWindow,self).__init__()
        uic.loadUi('finalwindow.ui',self)
        self.opendoc.clicked.connect(self.openDocument)
        self.show()
        
    def openDocument():
        global docpath  
        if docpath:
            doc_url = QUrl.fromLocalFile(docpath)
            QDesktopServices.openUrl(doc_url)
        else:
            QMessageBox.critical(None, "Error", "No document path available.")
        
        

        
        
                
def main():
    app = QApplication([])
    window = MyGUI()
    app.exec_()

if __name__ == '__main__':
    main()            